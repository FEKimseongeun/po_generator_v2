"""
HTML to Word Converter - Fixed Version
Handles missing styles gracefully
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from html.parser import HTMLParser
from io import StringIO


class HTMLToWordConverter(HTMLParser):
    """Parse HTML and convert to Word document elements"""

    def __init__(self, doc, parent_element=None, style_info=None):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        self.parent_element = parent_element

        # Style information to inherit from placeholder
        self.style_info = style_info or {
            'font_name': 'Arial',
            'font_size': Pt(11),
            'bold': False,
            'italic': False,
        }

        # State tracking
        self.in_bold = False
        self.in_italic = False
        self.in_list = False
        self.list_items = []
        self.in_table = False
        self.table_rows = []
        self.current_row = []
        self.in_header = False
        
    def handle_starttag(self, tag, attrs):
        """Handle opening HTML tags"""
        
        if tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()
            
        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run().add_break()
                
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Try to use heading style, fallback to bold paragraph
            level = int(tag[1])
            self.current_paragraph = self.doc.add_paragraph()
            
            try:
                # Try to apply heading style
                style_name = f'Heading {level}'
                self.current_paragraph.style = style_name
            except KeyError:
                # Style doesn't exist, use bold formatting instead
                self.current_paragraph.runs[0].bold = True if self.current_paragraph.runs else None
                # We'll apply bold in handle_data
                self.in_bold = True
            
        elif tag == 'strong' or tag == 'b':
            self.in_bold = True
            
        elif tag == 'em' or tag == 'i':
            self.in_italic = True
            
        elif tag == 'ol' or tag == 'ul':
            self.in_list = True
            self.list_items = []
            
        elif tag == 'li':
            pass
            
        elif tag == 'table':
            self.in_table = True
            self.table_rows = []
            
        elif tag == 'tr':
            self.current_row = []
            
        elif tag == 'th':
            self.in_header = True
            
        elif tag == 'td' or tag == 'th':
            pass
    
    def handle_endtag(self, tag):
        """Handle closing HTML tags"""
        
        if tag == 'p':
            self.current_paragraph = None
            
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.in_bold = False
            self.current_paragraph = None
            
        elif tag == 'strong' or tag == 'b':
            self.in_bold = False
            
        elif tag == 'em' or tag == 'i':
            self.in_italic = False
            
        elif tag == 'ol':
            # Add ordered list
            for i, item in enumerate(self.list_items, 1):
                p = self.doc.add_paragraph(f"{i}. {item}")
                p.paragraph_format.left_indent = Inches(0.5)
                # Apply inherited style
                for run in p.runs:
                    self._apply_style_to_run(run)
            self.in_list = False
            self.list_items = []

        elif tag == 'ul':
            # Add unordered list - use bullet character instead of style
            for item in self.list_items:
                p = self.doc.add_paragraph(f"• {item}")
                p.paragraph_format.left_indent = Inches(0.5)
                # Apply inherited style
                for run in p.runs:
                    self._apply_style_to_run(run)
            self.in_list = False
            self.list_items = []
            
        elif tag == 'tr':
            if self.current_row:
                self.table_rows.append(self.current_row)
            self.current_row = []
            
        elif tag == 'th':
            self.in_header = False
            
        elif tag == 'table':
            if self.table_rows:
                self._create_table()
            self.in_table = False
            self.table_rows = []
    
    def handle_data(self, data):
        """Handle text content"""
        data = data.strip()
        if not data:
            return

        if self.in_list:
            self.list_items.append(data)

        elif self.in_table:
            self.current_row.append(data)

        elif self.current_paragraph:
            run = self.current_paragraph.add_run(data)
            # Apply inherited style
            self._apply_style_to_run(run)
            # Apply inline formatting
            if self.in_bold:
                run.bold = True
            if self.in_italic:
                run.italic = True
        else:
            # Fallback: create paragraph if none exists
            self.current_paragraph = self.doc.add_paragraph()
            run = self.current_paragraph.add_run(data)
            # Apply inherited style
            self._apply_style_to_run(run)
            # Apply inline formatting
            if self.in_bold:
                run.bold = True
            if self.in_italic:
                run.italic = True

    def _apply_style_to_run(self, run):
        """Apply inherited style information to a run"""
        if self.style_info:
            if self.style_info.get('font_name'):
                run.font.name = self.style_info['font_name']
            if self.style_info.get('font_size'):
                run.font.size = self.style_info['font_size']
            # Don't apply bold/italic from style_info, those are for the placeholder itself
    
    def _create_table(self):
        """Create Word table from collected data with borders"""
        if not self.table_rows:
            return

        num_rows = len(self.table_rows)
        num_cols = max(len(row) for row in self.table_rows)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        # Set table style with borders
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Add table borders using XML
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Create table borders element
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Black border
            tblBorders.append(border)

        tblPr.append(tblBorders)

        # Fill table
        for row_idx, row_data in enumerate(self.table_rows):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_data

                # Apply inherited style to cell text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self._apply_style_to_run(run)
                        # Bold header row
                        if row_idx == 0:
                            run.bold = True


def html_to_word(html_content, doc, style_info=None):
    """
    Convert HTML string to Word document elements

    Args:
        html_content: HTML string to convert
        doc: python-docx Document object
        style_info: Optional dict with font_name, font_size, bold, italic
    """
    if not html_content or html_content == "null":
        return

    parser = HTMLToWordConverter(doc, style_info=style_info)

    try:
        parser.feed(html_content)
    except Exception as e:
        print(f"⚠️  Warning: HTML parsing error: {e}")
        # Fallback: add as plain text
        para = doc.add_paragraph(html_content)
        # Apply style to fallback paragraph
        if style_info and para.runs:
            for run in para.runs:
                if style_info.get('font_name'):
                    run.font.name = style_info['font_name']
                if style_info.get('font_size'):
                    run.font.size = style_info['font_size']