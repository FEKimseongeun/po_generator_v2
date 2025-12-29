"""
Word Document Generator - Complete Placeholder & Styling Solution

✅ FEATURES:
1. Inline Placeholder Replacement (multi-run support)
   - {{PO_NO}}-A01 → 210025-28-126-001-A01
   - {{DELIVERY_TERMS}} → actual delivery terms
   - Preserves original formatting (bold, italic, font, size)

2. HTML Block Replacement
   - {{PAYMENT_FULL}} → complete HTML content
   - {{WARRANTY}} → complete HTML content
   - Original placeholder completely removed

3. Style Inheritance
   - Extracts font name & size from placeholder location
   - Applies to all inserted HTML content (text, lists, tables)
   - Ensures consistent document styling

4. Table Borders
   - All HTML tables rendered with black borders
   - Headers automatically bolded
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import json
import os
import re
from html_to_word import html_to_word


class WordGenerator:
    """Generate Word document from HTML JSON data"""
    
    # Define which fields are simple text (inline) vs complex HTML (block)
    SIMPLE_FIELDS = ['PO_NO', 'MOM_DATE', 'DELIVERY_TERMS']
    
    def __init__(self, template_path=None):
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
        else:
            self.doc = Document()
            self._setup_default_styles()
        
        self._ensure_styles_exist()
    
    def _setup_default_styles(self):
        """Setup default document styles"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _ensure_styles_exist(self):
        """Ensure required styles exist"""
        for level in range(1, 7):
            style_name = f'Heading {level}'
            try:
                self.doc.styles[style_name]
            except KeyError:
                from docx.enum.style import WD_STYLE_TYPE
                style = self.doc.styles.add_style(
                    style_name, 
                    WD_STYLE_TYPE.PARAGRAPH
                )
                style.font.bold = True
                style.font.size = Pt(16 - level * 2)
    
    def generate_from_html_json(self, html_json_path, output_path):
        """Generate Word document without template"""
        print(f"📂 Loading JSON from: {html_json_path}")
        
        from json_validator import validate_and_fix_json_file
        data = validate_and_fix_json_file(html_json_path)
        html_data = data.get('html_data', {})
        
        if not html_data:
            raise ValueError("No 'html_data' found in JSON")
        
        print(f"✅ Loaded {len(html_data)} fields")
        
        self.doc.add_heading('Purchase Order', 0)
        self._add_header_info(html_data)
        
        self._add_section('Payment Terms', html_data.get('PAYMENT_FULL', ''))
        self._add_section('Warranty', html_data.get('WARRANTY', ''))
        self._add_section('Liquidated Damages', html_data.get('LIQUIDATED_DAMAGES', ''))
        self._add_section('Bond Requirements', html_data.get('BOND_FULL', ''))
        
        if html_data.get('DELIVERY_TERMS'):
            p = self.doc.add_paragraph()
            p.add_run('Delivery Terms: ').bold = True
            p.add_run(html_data.get('DELIVERY_TERMS', ''))
            self.doc.add_paragraph()
        
        self._add_section('Optional Items', html_data.get('OPTIONAL_FULL', ''))
        self._add_section('Supervision & Training', html_data.get('SUPERVISION_SHOP_TRAINING', ''))
        self._add_section('Special Notes', html_data.get('SPECIAL_NOTE', ''))
        self._add_section('Attachments', html_data.get('ATTACHMENT_FULL', ''))
        
        self.doc.save(output_path)
        print(f"✅ Word document saved: {output_path}")
        
        return output_path
    
    def _add_header_info(self, html_data):
        """Add PO number and date header"""
        p = self.doc.add_paragraph()
        p.add_run('PO No: ').bold = True
        p.add_run(html_data.get('PO_NO', 'N/A'))
        
        p = self.doc.add_paragraph()
        p.add_run('Date: ').bold = True
        p.add_run(html_data.get('MOM_DATE', 'N/A'))
        
        self.doc.add_paragraph()
    
    def _add_section(self, title, html_content, level=1):
        """Add a section to the document"""
        if not html_content or html_content == 'null':
            return
        
        self.doc.add_heading(title, level)
        
        try:
            html_to_word(html_content, self.doc)
        except Exception as e:
            print(f"⚠️  Warning: Error converting HTML for {title}: {e}")
            self.doc.add_paragraph(html_content)
        
        self.doc.add_paragraph()
    
    def generate_with_template(self, template_path, html_json_path, output_path):
        """
        Generate Word using template with {{placeholders}}
        PROPERLY handles inline replacement and HTML conversion
        """
        print(f"📂 Loading template: {template_path}")
        self.doc = Document(template_path)
        
        print(f"📂 Loading JSON: {html_json_path}")
        from json_validator import validate_and_fix_json_file
        data = validate_and_fix_json_file(html_json_path)
        html_data = data.get('html_data', {})
        
        print(f"✅ Loaded {len(html_data)} fields")
        
        # Process all paragraphs
        print("🔄 Processing paragraphs...")
        paragraphs_to_process = []
        
        for para in self.doc.paragraphs:
            paragraphs_to_process.append(para)
        
        for para in paragraphs_to_process:
            self._process_paragraph(para, html_data)
        
        # Process tables
        print("🔄 Processing tables...")
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._process_paragraph(para, html_data)
        
        # Save
        self.doc.save(output_path)
        print(f"✅ Word document saved: {output_path}")
        
        return output_path
    
    def _process_paragraph(self, paragraph, html_data):
        """
        Process a single paragraph, replacing placeholders

        Strategy:
        1. Simple fields (PO_NO, MOM_DATE, DELIVERY_TERMS):
           - Inline replacement preserving original text formatting
           - Example: "PO No. {{PO_NO}}-A01" → "PO No. 210025-28-126-001-A01"

        2. Complex fields (PAYMENT_FULL, WARRANTY, etc.):
           - Replace entire paragraph with HTML-rendered content
           - Inherit font/size from original placeholder location
        """

        full_text = paragraph.text

        # Check if paragraph contains any placeholders
        placeholders_found = re.findall(r'\{\{([A-Z_]+)\}\}', full_text)

        if not placeholders_found:
            return

        print(f"   📌 Found placeholders: {placeholders_found}")
        print(f"      In text: {full_text[:80]}...")

        # Separate simple vs complex placeholders
        simple_placeholders = [p for p in placeholders_found if p in self.SIMPLE_FIELDS]
        complex_placeholders = [p for p in placeholders_found if p not in self.SIMPLE_FIELDS]

        if simple_placeholders:
            print(f"      🔤 Inline replacement for: {simple_placeholders}")
        if complex_placeholders:
            print(f"      📄 HTML replacement for: {complex_placeholders}")

        # Handle simple placeholders - INLINE REPLACEMENT
        if simple_placeholders:
            self._replace_simple_placeholders(paragraph, html_data, simple_placeholders)

        # Handle complex placeholders - INSERT HTML BLOCKS AFTER
        if complex_placeholders:
            self._insert_complex_placeholders(paragraph, html_data, complex_placeholders)
    
    def _replace_simple_placeholders(self, paragraph, html_data, placeholders):
        """
        Replace simple placeholders inline (e.g., {{PO_NO}}-A01 → 210025-28-126-001-A01)

        CRITICAL: Word documents often split text across multiple runs, so "{{PO_NO}}"
        might be stored as "{{PO_" + "NO}}" in separate runs. We must process the
        entire paragraph text as a whole, not run-by-run.
        """

        # Get full paragraph text (concatenate all runs)
        full_text = paragraph.text

        # Build replacement map
        replacements = {}
        for placeholder in placeholders:
            value = html_data.get(placeholder, f"[{placeholder} NOT FOUND]")
            placeholder_tag = f"{{{{{placeholder}}}}}"
            replacements[placeholder_tag] = value

        # Replace all placeholders in the full text
        modified_text = full_text
        for placeholder_tag, value in replacements.items():
            if placeholder_tag in modified_text:
                modified_text = modified_text.replace(placeholder_tag, value)
                print(f"         ✓ {placeholder_tag} → {value}")

        # If text changed, rebuild the paragraph runs
        if modified_text != full_text:
            # Save the formatting of the first run (if exists)
            first_run_format = None
            if paragraph.runs:
                first_run = paragraph.runs[0]
                first_run_format = {
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                    'underline': first_run.underline,
                    'font_name': first_run.font.name,
                    'font_size': first_run.font.size,
                }

            # Clear all runs
            for run in paragraph.runs:
                run.text = ""

            # Add new run with replaced text
            new_run = paragraph.add_run(modified_text)

            # Restore formatting
            if first_run_format:
                if first_run_format['bold'] is not None:
                    new_run.bold = first_run_format['bold']
                if first_run_format['italic'] is not None:
                    new_run.italic = first_run_format['italic']
                if first_run_format['underline'] is not None:
                    new_run.underline = first_run_format['underline']
                if first_run_format['font_name']:
                    new_run.font.name = first_run_format['font_name']
                if first_run_format['font_size']:
                    new_run.font.size = first_run_format['font_size']
    
    def _insert_complex_placeholders(self, paragraph, html_data, placeholders):
        """
        Insert complex HTML content after the placeholder paragraph
        Extracts formatting from placeholder and applies to inserted HTML content
        """

        # Extract style information from the placeholder paragraph
        style_info = self._extract_paragraph_style(paragraph)

        # Get parent and position
        parent = paragraph._element.getparent()
        placeholder_index = parent.index(paragraph._element)

        # Get full paragraph text to check for placeholders
        full_text = paragraph.text

        # Check if paragraph contains ONLY placeholders (should be removed)
        # or has other text (should keep non-placeholder parts)
        has_other_text = False
        temp_text = full_text
        for placeholder in placeholders:
            placeholder_tag = f"{{{{{placeholder}}}}}"
            temp_text = temp_text.replace(placeholder_tag, "").strip()

        if temp_text:  # Has other text besides placeholders
            has_other_text = True

        # Remove placeholder paragraph completely (will be replaced with HTML content)
        parent.remove(paragraph._element)

        # Insert HTML content at the position
        for placeholder in placeholders:
            html_content = html_data.get(placeholder, '')

            if html_content and html_content != 'null':
                print(f"         ✓ Inserting HTML for {{{{{placeholder}}}}} with style: {style_info['font_name']}, {style_info['font_size']}")

                # Create a temporary document to render HTML
                temp_doc = Document()
                html_to_word(html_content, temp_doc, style_info)

                # Copy all elements from temp doc to main doc at position
                for element in temp_doc.element.body:
                    # Insert after placeholder position
                    parent.insert(placeholder_index, element)
                    placeholder_index += 1

    def _extract_paragraph_style(self, paragraph):
        """
        Extract style information from a paragraph (font, size, etc.)
        """
        style_info = {
            'font_name': None,
            'font_size': None,
            'bold': False,
            'italic': False,
        }

        # Get style from first run if available
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.font.name:
                style_info['font_name'] = first_run.font.name
            if first_run.font.size:
                style_info['font_size'] = first_run.font.size
            if first_run.bold:
                style_info['bold'] = first_run.bold
            if first_run.italic:
                style_info['italic'] = first_run.italic

        # Fallback to paragraph style
        if not style_info['font_name'] or not style_info['font_size']:
            try:
                para_style = paragraph.style
                if para_style and para_style.font:
                    if not style_info['font_name'] and para_style.font.name:
                        style_info['font_name'] = para_style.font.name
                    if not style_info['font_size'] and para_style.font.size:
                        style_info['font_size'] = para_style.font.size
            except:
                pass

        # Ultimate fallback
        if not style_info['font_name']:
            style_info['font_name'] = 'Arial'
        if not style_info['font_size']:
            style_info['font_size'] = Pt(11)

        return style_info


# CLI Usage
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python word_generator.py <html_json_path> <output_path> [template_path]")
        sys.exit(1)
    
    html_json_path = sys.argv[1]
    output_path = sys.argv[2]
    template_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    try:
        generator = WordGenerator()
        
        if template_path:
            print("📝 Using template mode")
            generator.generate_with_template(template_path, html_json_path, output_path)
        else:
            print("📝 Using default mode")
            generator.generate_from_html_json(html_json_path, output_path)
        
        print("\n✅ Success!")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)